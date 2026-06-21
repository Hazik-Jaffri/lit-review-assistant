"""
report_export.py
------------------
Exports the comparative analysis table to formats the student can drop
straight into their research report:

    - CSV   : for quick spreadsheet editing
    - DOCX  : a ready-to-paste Word table matching the exact 5-column
              schema required by the course guidelines (Reference, Title,
              Proposed Solution, Limitations/Gaps, Our Contribution), with
              the Confidence column included as a 6th column so the
              student can see which rows need manual review before
              deleting that column for the final submission.

DOCX generation uses the `python-docx` library (pure Python, no Node.js /
npm dependency), so it works the same way on the student's own machine,
Streamlit Community Cloud, or any other deployment target.
"""

from __future__ import annotations

import csv
import io

from . import config


def export_csv_bytes(table_rows, columns=None):
    """Return CSV file content as bytes, ready for st.download_button."""
    cols = columns or config.TABLE_COLUMNS
    buffer = io.StringIO()
    writer = csv.DictWriter(buffer, fieldnames=cols, extrasaction="ignore")
    writer.writeheader()
    for row in table_rows:
        writer.writerow(row)
    return buffer.getvalue().encode("utf-8")


# Relative column widths (inches) for a landscape US Letter page with 0.6"
# margins on each side (content width = 11 - 1.2 = 9.8").
_COLUMN_WIDTH_PLAN_IN = {
    "Reference": 1.0,
    "Title": 1.7,
    "Proposed Solution": 2.4,
    "Limitations / Gaps": 2.4,
    "Our Contribution": 1.9,
    "Confidence": 0.8,
}


def export_docx_bytes(table_rows, columns=None, title="Comparative Analysis Table", subtitle=""):
    """
    Render the table as a .docx file using python-docx and return the
    file's bytes. Raises RuntimeError with a clear message if python-docx
    is unavailable, so the UI can fall back to CSV.
    """
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "The 'python-docx' package is not installed. Run: pip install python-docx"
        ) from exc

    cols = columns or config.TABLE_COLUMNS

    document = Document()

    # Landscape US Letter with narrower margins so the table has room to breathe.
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Inches(0.6)
    section.left_margin = section.right_margin = Inches(0.6)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    heading = document.add_heading(title or "Comparative Analysis Table", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    if subtitle:
        sub = document.add_paragraph()
        sub_run = sub.add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(9)

    content_width_in = (section.page_width.inches) - 1.2  # minus left+right margins
    raw_widths = [_COLUMN_WIDTH_PLAN_IN.get(c, content_width_in / len(cols)) for c in cols]
    scale = content_width_in / sum(raw_widths)
    col_widths = [Inches(w * scale) for w in raw_widths]

    table = document.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    def _set_cell_shading(cell, hex_color):
        shd = cell._tc.get_or_add_tcPr()
        shading_elm = shd.makeelement(
            qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color}
        )
        shd.append(shading_elm)

    def _set_cell_width(cell, width):
        cell.width = width

    # Header row
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(cols):
        cell = header_cells[i]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(col_name)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_shading(cell, "1F4E79")
        _set_cell_width(cell, col_widths[i])

    # Body rows
    for row_data in table_rows:
        row_cells = table.add_row().cells
        for i, col_name in enumerate(cols):
            cell = row_cells[i]
            value = row_data.get(col_name, "")
            cell.text = "" if value is None else str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(9) if cell.paragraphs[0].runs else None
            _set_cell_width(cell, col_widths[i])

    # python-docx needs widths set on every cell in every row for tables to
    # render consistently in Word (table.autofit=False relies on this).
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
